import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()
    
    # ---------------------------------------------------------
    # PAGE SETUP & MARGINS
    # ---------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # ---------------------------------------------------------
    # COLOR PALETTE (PREMIUM DARK EXECUTIVE / SLATE & TEAL)
    # ---------------------------------------------------------
    COLOR_NAVY = RGBColor(15, 23, 42)      # #0F172A Slate 900
    COLOR_TEAL = RGBColor(13, 148, 136)    # #0D9488 Teal 600
    COLOR_BLUE = RGBColor(2, 132, 199)     # #0284C7 Sky 600
    COLOR_SLATE_700 = RGBColor(51, 65, 85) # #334155
    COLOR_SLATE_500 = RGBColor(100, 116, 139) # #64748B
    
    HEX_PRIMARY = "0F172A"
    HEX_TEAL = "0D9488"
    HEX_BLUE = "0284C7"
    HEX_LIGHT_BG = "F8FAFC"
    HEX_CODE_BG = "F1F5F9"
    HEX_BORDER = "CBD5E1"
    HEX_CALLOUT_INFO_BG = "F0F9FF"
    HEX_CALLOUT_INFO_BORDER = "0284C7"
    HEX_CALLOUT_WARN_BG = "FFFBEB"
    HEX_CALLOUT_WARN_BORDER = "F59E0B"
    
    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_SLATE_700
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)
    
    # ---------------------------------------------------------
    # HELPER FUNCTIONS
    # ---------------------------------------------------------
    def set_cell_background(cell, hex_color):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, size in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            m = OxmlElement(side)
            m.set(qn('w:w'), str(size))
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

    def set_table_borders(table, hex_color="CBD5E1"):
        tblPr = table._tbl.tblPr
        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{hex_color}"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
        '''
        tblPr.append(parse_xml(borders_xml))

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_TEAL
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_BLUE
        return p

    def add_p(text="", bold_prefix=None, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_NAVY
        if text:
            r_text = p.add_run(text)
            r_text.font.color.rgb = COLOR_SLATE_700
        return p

    def add_bullet(bold_prefix, text=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_NAVY
        if text:
            r_text = p.add_run(text)
            r_text.font.color.rgb = COLOR_SLATE_700
        return p

    def add_callout(title, text, style_type="info"):
        bg_color = HEX_CALLOUT_INFO_BG if style_type == "info" else HEX_CALLOUT_WARN_BG
        border_color = HEX_CALLOUT_INFO_BORDER if style_type == "info" else HEX_CALLOUT_WARN_BORDER
        
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders_xml = f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
        '''
        tcPr.append(parse_xml(borders_xml))
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(f"📌 {title}\n")
        r_title.font.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = COLOR_BLUE if style_type == "info" else RGBColor(217, 119, 6)
        
        r_text = p.add_run(text)
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = COLOR_SLATE_700
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, HEX_CODE_BG)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders_xml = f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:color="CBD5E1"/>
            <w:left w:val="single" w:sz="12" w:color="0284C7"/>
            <w:bottom w:val="single" w:sz="4" w:color="CBD5E1"/>
            <w:right w:val="single" w:sz="4" w:color="CBD5E1"/>
        </w:tcBorders>
        '''
        tcPr.append(parse_xml(borders_xml))
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_NAVY
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_styled_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        
        # Header Row
        hdr_row = table.rows[0]
        hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for i, header_text in enumerate(headers):
            cell = hdr_row.cells[i]
            if col_widths and i < len(col_widths):
                cell.width = Inches(col_widths[i])
            set_cell_background(cell, HEX_PRIMARY)
            set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
        # Data Rows
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            bg_color = HEX_LIGHT_BG if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                if col_widths and c_idx < len(col_widths):
                    cell.width = Inches(col_widths[c_idx])
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(val))
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_SLATE_700

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(30)
    p_logo.paragraph_format.space_after = Pt(20)
    
    logo_path = os.path.join('public', 'fluxo-logo-v2.png')
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.5))
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(8)
    r_t = p_title.add_run("DOCUMENTAÇÃO TÉCNICA E ARQUITETURAL")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(26)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_NAVY
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run("SISTEMA FLUXO — GESTÃO FINANCEIRA INTELIGENTE\n"
                          "Especificação Completa de Engenharia, Banco de Dados, Módulos e Guia de Implantação")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = COLOR_TEAL

    # Meta Table on Cover Page
    add_styled_table(
        ["Atributo do Documento", "Detalhamento de Engenharia / Transferência"],
        [
            ["Nome do Produto", "Fluxo Financeiro (Web & Mobile PWA)"],
            ["Versão do Sistema", "1.0.0 (Homologado e Pronto para Produção)"],
            ["Data de Emissão", "11 de Agosto de 2026"],
            ["Autor / Responsável", "Analista de Sistemas & Engenheiro de Software Lead"],
            ["Finalidade do Documento", "Dossiê Técnico Completo para Transferência de Propriedade Intelectual (Handoff de Venda Comercial)"],
            ["Stack Principal", "React 18, Vite 5, TypeScript 5, Supabase (PostgreSQL 15), Tailwind CSS, PWA"],
            ["Status do Projeto", "100% Funcional, Testes Automatizados Validados, RLS Ativado"]
        ],
        col_widths=[2.2, 4.3]
    )

    doc.add_page_break()

    # ---------------------------------------------------------
    # HEADERS & FOOTERS SETUP FOR SUBSEQUENT PAGES
    # ---------------------------------------------------------
    section = doc.sections[0]
    header = section.header
    p_hdr = header.paragraphs[0]
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_hdr = p_hdr.add_run("FLUXO FINANCEIRO — DOCUMENTAÇÃO TÉCNICA DE SISTEMAS")
    r_hdr.font.size = Pt(8.5)
    r_hdr.font.color.rgb = COLOR_SLATE_500
    
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ftr = p_ftr.add_run("Fluxo Financeiro © 2026 — Documento de Handoff Técnico | Confidencial")
    r_ftr.font.size = Pt(8.5)
    r_ftr.font.color.rgb = COLOR_SLATE_500

    # ---------------------------------------------------------
    # 1. SUMÁRIO EXECUTIVO E VISÃO GERAL DO PRODUTO
    # ---------------------------------------------------------
    add_h1("1. Sumário Executivo e Visão Geral do Produto")
    
    add_p("O ", "O Sistema Fluxo Financeiro ")
    p_desc = doc.paragraphs[-1]
    p_desc.add_run("é uma plataforma completa de gestão financeira pessoal e empresarial desenvolvida para oferecer controle de caixa, inteligência de planejamento, gestão de passivos (dívidas) e projeções patrimoniais com alta performance e máxima fluidez de uso. O produto foi concebido sob princípios modernos de engenharia de software para operar como um aplicativo web e PWA (Progressive Web App) responsivo, multiplataforma e multi-tenant.")

    add_h2("1.1. Propósito e Valor de Negócio")
    add_p("Diferente de planilhas tradicionais ou aplicativos financeiros excessivamente complexos, o Fluxo resolve a dor central do controle financeiro doméstico e de pequenas empresas: a desorganização de fluxo de caixa e o descontrole de faturas de cartão de crédito. Ele fornece uma visão executiva direta, permitindo saber exatamente o saldo consolidado atual, despesas previstas vs. efetivadas, capacidade de poupança e score de adimplência.")

    add_h2("1.2. A Filosofia de Negócio: Regime de Caixa (Visão Extrato)")
    add_p("Uma das decisões arquiteturais e de regras de negócio mais importantes do sistema é a aderência estrita ao ", "Regime de Caixa para Liquidação de Lançamentos: ")
    add_bullet("Comportamento de Carteira/Extrato: ", "O aplicativo contabiliza e exibe despesas e receitas normais (não-cartão) com base estrita na sua Data de Baixa/Pagamento (payment_date). A data de vencimento/nominal serve como referência de agendamento, mas os saldos reais e extratos refletem exatamente quando o dinheiro entrou ou saiu da conta.")
    add_bullet("Faturas de Cartão de Crédito: ", "O valor de faturas de cartão de crédito e suas parcelas é baixado e contabilizado no extrato bancário no momento em que a fatura do mês correspondente (invoiceMonthYear) é quitada pelo usuário.")
    add_bullet("Transferências entre Contas: ", "Transferências não alteram a receita nem a despesa líquida do patrimônio global do usuário; elas movimentam o saldo entre contas de origem e destino usando um identificador de vínculo relacional (transfer_group_id).")

    add_h2("1.3. Diretrizes de UX/UI e Design System")
    add_p("O aplicativo adota um padrão de design visual profissional e executivo (Apple Minimalist Standard) projetado para transmitir confiabilidade bancária:")
    add_bullet("Leitura Executiva: ", "Prioridade para métricas numéricas comparativas, cards objetivos e status visuais em detrimento de textos onboarding expositivos longos.")
    add_bullet("Proibição de Emojis em UI: ", "A interface utiliza exclusivamente iconografia SVG vetorial padronizada (Lucide React) para manter a máxima sobriedade visual e profissionalismo.")
    add_bullet("Tema Escuro Premium & Suporte a Acessibilidade: ", "Interface baseada em tokens CSS HSL com alto contraste, suporte a modo claro/escuro e personalização dinâmica da cor de destaque (Theme Accent Color) por usuário.")

    add_callout(
        "Princípio do Produto — Menos Explicação, Mais Ação",
        "Toda a complexidade de cálculos (projeção de juros, score aditivo, regras de conciliação) fica encapsulada no motor do backend e em hooks de domínio. A interface exibe apenas indicadores limpos e acionáveis com tooltips discretos quando necessário.",
        "info"
    )

    # ---------------------------------------------------------
    # 2. ARQUITETURA DE SOFTWARE E PILHA TECNOLÓGICA
    # ---------------------------------------------------------
    add_h1("2. Arquitetura de Software e Pilha Tecnológica")

    add_p("O Fluxo foi projetado utilizando a arquitetura ", "BaaS (Backend-as-a-Service) desacoplada com Single Page Application (SPA) reativa e PWA.")

    add_h2("2.1. Análise de Paradigmas de Programação: Orientação a Objetos (POO) vs Programação Funcional (FP)")
    add_p("Ao analisar a base de código do Fluxo como Analista de Sistemas, observa-se uma abordagem híbrida altamente eficiente e arquiteturalmente limpa:")
    
    add_bullet("Camada de Apresentação e Hooks (Programação Funcional): ", "O frontend em React 18 utiliza o paradigma Funcional puro. Componentes são funções puras com renderização declarativa, imutabilidade de estado gerenciada por hooks (useState, useMemo, useCallback) e gerenciamento de estado assíncrono via TanStack Query v5 e Zustand.")
    add_bullet("Camada de Domínio, Serviços e Abstração de Dados (Orientação a Objetos): ", "O modelo de domínio do sistema implementa rigorosamente os princípios de POO:")
    add_bullet("  • Encapsulamento & Responsabilidade Única: ", "Serviços dedicados (ex: transactionService.ts) encapsulam a lógica de comunicação externa, transformação de payload e tratamento de exceções, isolando os componentes visuais dos detalhes do banco de dados.")
    add_bullet("  • Contratos & Interfaces Fortes: ", "O TypeScript é utilizado para definir contratos de dados estritos (Types e Interfaces em src/types/finance.ts), garantindo tipagem segura em tempo de compilação para entidades como Transaction, Account, CreditCard, Debt e Budget.")
    add_bullet("  • Encapsulamento em Banco de Dados (Procedural & Triggers OO): ", "O PostgreSQL no Supabase armazena regras de negócios críticas através de Triggers e Stored Procedures (RPCs como delete_user_data e update_account_balance_trigger), comportando-se como métodos de classe encapsulados na camada de dados.")

    add_h2("2.2. Detalhamento da Stack Tecnológica")
    
    add_styled_table(
        ["Camada", "Tecnologia / Bibliotecas", "Versão", "Função Arquitetural no Sistema"],
        [
            ["Core Frontend", "React", "18.3.1", "Biblioteca UI declarativa baseada em componentes reativos."],
            ["Linguagem", "TypeScript", "5.8.3", "Tipagem estática rigorosa para segurança e prevenção de erros."],
            ["Build Tool", "Vite (com SWC)", "5.4.19", "Bundler ultra-rápido com compilação SWC e Hot Module Replacement."],
            ["Estilização", "Tailwind CSS + PostCSS", "3.4.17", "Framework CSS utilitário com sistema de design tokens dinâmicos."],
            ["Componentes UI", "Radix UI / shadcn/ui", "1.x / 2.x", "Primitivas de interface acessíveis (Dialogs, Popovers, Selects, Tabs)."],
            ["Gerenciamento Estado", "Zustand", "5.0.11", "Store global leve para preferências de usuário, temas e estados temporários."],
            ["Gerenciamento Servidor", "TanStack Query", "5.95.0", "Cache inteligente, revalidação em background e mutações otimistas."],
            ["Formulários e Validação", "React Hook Form + Zod", "7.61 / 3.25", "Gerenciamento de formulários de alta performance com schema validation."],
            ["Gráficos e Visualização", "Recharts", "2.15.4", "Visualização de dados financeiros (gráficos de pizza, área e barras)."],
            ["Iconografia", "Lucide React", "0.462.0", "Conjunto de ícones vetoriais padronizados para UI."],
            ["Backend as a Service", "Supabase (PostgreSQL)", "2.95.1", "Banco de dados relacional, autenticação JWT, RLS e Storage."],
            ["Edge Serverless", "Supabase Edge Functions", "Deno 2.x", "Execução de rotinas serverless protegidas (ex: envio de Notificações Push)."],
            ["Progressive Web App", "Vite PWA / Service Worker", "1.2.0", "Suporte a instalação mobile/desktop, offline-ready e notificações push."],
            ["Testes e Qualidade", "Vitest + React Testing Library", "3.2.4", "Suíte de testes automatizados unitários e de integração."]
        ],
        col_widths=[1.3, 1.6, 0.8, 2.8]
    )

    # ---------------------------------------------------------
    # 3. MODELAGEM DE BANCO DE DADOS E ENGENHARIA DE DADOS
    # ---------------------------------------------------------
    add_h1("3. Modelagem de Banco de Dados e Engenharia de Dados")

    add_p("O banco de dados do Fluxo é estruturado sobre o ", "PostgreSQL 15+ hospedado na infraestrutura Supabase, ")
    p_db = doc.paragraphs[-1]
    p_db.add_run("com arquitetura fortemente relacional, chaves estrangeiras com suporte a ON DELETE CASCADE/SET NULL, índices de escalabilidade e isolamento total multi-tenant garantido por Row Level Security (RLS).")

    add_h2("3.1. Dicionário de Tabelas Principais")

    add_styled_table(
        ["Tabela SQL", "Campos Chave / Tipos", "Descrição e Regras de Integridade"],
        [
            [
                "users_profile / profiles",
                "id (UUID, PK, FK auth.users)\ncurrency (TEXT)\ntheme_color (TEXT)\nplan_id (UUID, FK plans)\ncreated_at (TIMESTAMPTZ)",
                "Armazena dados de perfil do usuário, moeda preferida, cor de destaque configurada no app e vínculo de plano de licença."
            ],
            [
                "accounts",
                "id (UUID, PK)\nuser_id (UUID, FK)\nname (TEXT)\ntype (TEXT)\nbalance (NUMERIC(15,2))\ninitial_balance (NUMERIC)\noverdraft_limit (NUMERIC)\ndeleted_at (TIMESTAMPTZ)",
                "Contas bancárias/carteiras do usuário. O campo balance é mantido automaticamente por triggers SQL. Suporta soft delete via deleted_at."
            ],
            [
                "credit_cards",
                "id (UUID, PK)\nuser_id (UUID, FK)\nname (TEXT)\nlimit_amount (NUMERIC)\nclosing_day (INT)\ndue_day (INT)\ncolor (TEXT)",
                "Cartões de crédito do usuário. Armazena o limite concedido e as datas de corte (fechamento) e vencimento da fatura."
            ],
            [
                "categories",
                "id (UUID, PK)\nuser_id (UUID, FK)\nname (TEXT)\ntype (TEXT - expense/income)\ncategory_group_id (UUID, FK)\nbudget_limit (NUMERIC)\nis_fixed (BOOLEAN)",
                "Categorias e subcategorias financeiras. Permite definir se é despesa fixa, teto de orçamento mensal e pertencimento a agrupamentos."
            ],
            [
                "transactions",
                "id (UUID, PK)\nuser_id (UUID, FK)\ndescription (TEXT)\namount (NUMERIC(15,2))\ntype (TEXT)\ncategory_id (UUID, FK)\naccount_id (UUID, FK)\ncard_id (UUID, FK)\ndate (DATE)\npayment_date (DATE)\nis_paid (BOOLEAN)\nis_transfer (BOOLEAN)\ntransfer_group_id (UUID)\ninstallment_number (INT)\ntotal_installments (INT)\ndeleted_at (TIMESTAMPTZ)",
                "Tabela central de lançamentos. Suporta movimentações em conta, despesas no cartão, parcelamentos, agendamentos e transferências casadas (transfer_group_id)."
            ],
            [
                "bills",
                "id (UUID, PK)\nuser_id (UUID, FK)\ndescription (TEXT)\namount (NUMERIC)\ndue_date (DATE)\nstatus (TEXT)\ncard_id (UUID, FK)\npaid_at (TIMESTAMPTZ)",
                "Contas a pagar/receber agrupadas e faturas registradas do cartão de crédito."
            ],
            [
                "debts",
                "id (UUID, PK)\nuser_id (UUID, FK)\ncreditor_name (TEXT)\ntotal_amount (NUMERIC)\nstatus (TEXT)\ninterest_rate (NUMERIC)\noriginal_debt_id (UUID)",
                "Registro de passivos (dívidas) e acordos de renegociação financeira com histórico de amortização."
            ],
            [
                "goals / emergency_funds",
                "id (UUID, PK)\nuser_id (UUID, FK)\ntarget_amount (NUMERIC)\ncurrent_amount (NUMERIC)\nmonthly_yield_rate (NUMERIC)\ntarget_date (DATE)",
                "Metas financeiras e configuração da reserva de emergência (custo de vida mensal, número de meses e taxa de rendimento mensal)."
            ],
            [
                "plans / global_flags / super_admins",
                "id (UUID, PK)\ncode (TEXT - basic/pro/premium)\nlimits_json (JSONB)\nis_active (BOOLEAN)\nflag_name (TEXT)",
                "Tabelas de governança do sistema para gestão de limites de plano, liberação de funcionalidades e controle do Painel Super Admin."
            ]
        ],
        col_widths=[1.5, 2.2, 2.8]
    )

    add_h2("3.2. Triggers Automatizados e Regras de Banco de Dados")
    add_p("O sistema não confia apenas na lógica do cliente para manter a integridade dos dados. Ele utiliza triggers PostgreSQL resilientes:")
    
    add_bullet("Trigger de Atualização de Saldo Bancário (update_account_balance_trigger): ", "Sempre que uma transação é inserida, alterada ou excluída (respeitando soft delete), este trigger recalcula automaticamente o saldo atual (balance) da conta correspondente somando movimentações efetivadas (is_paid = true).")
    add_bullet("Trigger de Lançamentos Recorrentes (trigger_spawn_recurring): ", "Responsável por gerar automaticamente novos lançamentos futuros para despesas e receitas marcadas como recorrentes.")
    add_bullet("RPC de Exclusão de Usuário e LGPD (delete_user_data): ", "Stored procedure segura executada com privilégios de SECURITY DEFINER que limpa em lote todos os dados de transações, contas, cartões, categorias e dívidas quando um usuário solicita o encerramento da conta, garantindo conformidade com a LGPD.")

    add_h2("3.3. Políticas de Segurança (Row Level Security - RLS)")
    add_p("Todas as tabelas do banco de dados possuem políticas RLS ativas. O acesso é estritamente isolado pelo UUID do usuário autenticado no JWT:")
    add_code_block(
        "-- Exemplo de Política RLS aplicada na tabela de transações\n"
        "CREATE POLICY \"Usuários acessam apenas suas próprias transações\"\n"
        "ON public.transactions\n"
        "FOR ALL\n"
        "USING (auth.uid() = user_id)\n"
        "WITH CHECK (auth.uid() = user_id);"
    )

    # ---------------------------------------------------------
    # 4. ESPECIFICAÇÃO DE MÓDULOS E REGRAS DE NEGÓCIO
    # ---------------------------------------------------------
    doc.add_page_break()
    add_h1("4. Especificação de Módulos e Regras de Negócio")
    add_p("A seguir, detalha-se o funcionamento técnico e de negócio de cada módulo integrante do Fluxo Financeiro, explicando ", "o que cada recurso faz e como faz.")

    add_h2("4.1. Módulo de Autenticação e Gestão de Acesso")
    add_bullet("O que faz: ", "Gerencia a entrada segura do usuário, redefinição de senha, autenticação via Magic Link e gestão de sessão.")
    add_bullet("Como faz: ", "Utiliza o Supabase Auth com persistência de token JWT em LocalStorage. Possui páginas dedicadas para fluxo de recuperação de senha com redirecionamento automático (AuthPage.tsx, EmailResetPasswordPage.tsx, ReauthenticationPage.tsx) e tratamento seguro contra race conditions de sessão.")

    add_h2("4.2. Módulo de Dashboard & Boot Inteligente")
    add_bullet("O que faz: ", "Exibe a visão geral consolidada da saúde financeira do usuário no mês selecionado.")
    add_bullet("Como faz: ", "Utiliza componentes minimalistas (Metric Cards estilo Apple). Ao abrir o aplicativo, o motor de Boot Inteligente bloqueia a exibição de dados zerados falsos enquanto executa o refresh real com o Supabase. Se houver falha de rede, um fallback gracioso carrega os últimos dados salvos.")

    add_h2("4.3. Módulo de Gestão de Contas Bancárias & Extrato")
    add_bullet("O que faz: ", "Permite cadastrar carteiras, contas correntes, poupanças e investimentos, ajustando saldos e registrando conciliações.")
    add_bullet("Como faz: ", "Integra o useAccountMutations.ts. Suporta configuração de limite de cheque especial (overdraft_limit) e marcação de contas inativas via soft delete (deleted_at). O extrato reflete a data de baixa efetiva.")

    add_h2("4.4. Módulo de Cartões de Crédito & Gestão de Faturas")
    add_bullet("O que faz: ", "Controla o limite de cartões de crédito, calcula o fechamento e vencimento de faturas, realiza parcelamento automático e permite abatimentos parciais ou totais.")
    add_bullet("Como faz: ", "Calcula dinamicamente a competência da fatura (invoiceMonthYear) com base no dia de fechamento. Ao pagar a fatura, o sistema solicita a conta de origem bancária para realizar o débito real ou permite registrar o abatimento sem afetar o caixa caso tenha sido pago por terceiros. Oferece suporte a Pix no Crédito / Transferência via Cartão.")

    add_h2("4.5. Módulo de Lançamentos & Transferências entre Contas")
    add_bullet("O que faz: ", "Permite incluir receitas, despesas (únicas, fixas ou parceladas) e realizar transferências entre contas.")
    add_bullet("Como faz: ", "Nas transferências, o sistema cria dois lançamentos vinculados por um transfer_group_id (uma saída na conta de origem e uma entrada na conta de destino), rotulados com a categoria nativa de Transferência, garantindo que o valor não infle as receitas ou despesas nos relatórios.")

    add_h2("4.6. Módulo de Orçamentos & Limites de Gastos")
    add_bullet("O que faz: ", "Define teto orçamentário por categoria ou agrupamento e acompanha o percentual de consumo em tempo real.")
    add_bullet("Como faz: ", "O hook useBudgetGroups.ts compara os lançamentos efetivados e previstos do mês com o campo budget_limit da categoria, exibindo barras comparativas e alertas visuais de estouro de orçamento.")

    add_h2("4.7. Módulo de Dívidas & Acordos de Renegociação")
    add_bullet("O que faz: ", "Gerencia débitos com credores, permitindo registrar acordos de quitação com valor de entrada e parcelamento do saldo remanescente.")
    add_bullet("Como faz: ", "O hook useDebtMutations.ts converte dívidas em acordos ativos, gerando os lançamentos de entrada no caixa e acompanhando a amortização gradual até a quitação total do passivo.")

    add_h2("4.8. Módulo de Reserva de Emergência & Simulação de Rendimentos")
    add_bullet("O que faz: ", "Calcula o tamanho ideal da reserva financeira com base no custo de vida mensal e simula a evolução com juros compostos.")
    add_bullet("Como faz: ", "O hook useEmergencyFund.ts multiplica a média de despesas mensais pelo número de meses desejado (ex: 6 meses). A simulação aplica a taxa mensal cadastrada (monthly_yield_rate) para calcular a projeção patrimonial.")

    add_h2("4.9. Módulo \"Fluxo Score\" (Inteligência Financeira Algorítmica)")
    add_bullet("O que faz: ", "Avalia a saúde financeira do usuário fornecendo uma pontuação de 0 a 1000 acompanhada de diagnóstico executivo.")
    add_bullet("Como faz: ", "Algoritmo aditivo e puramente de leitura que analisa a pontualidade do pagamento de contas de consumo, a existência de dívidas atrasadas, a relação despesa/receita e concede bônus no primeiro dia útil do mês. A UI exibe um medidor gráfico circular (Gauge) animado.")

    add_h2("4.10. Módulo de Relatórios Gerenciais & Exportação em PDF")
    add_bullet("O que faz: ", "Fornece análises de consumo por banco, categoria e subcategoria, com emissão de relatório gerencial em PDF e layout de impressão de alta fidelidade.")
    add_bullet("Como faz: ", "ReportsDashboard.tsx processa cruzamentos de dados via Recharts e gera modal de impressão com CSS paged media (@media print) estilizado sem quebras de página inadequadas.")

    add_h2("4.11. Módulo Super Admin & Gestão de Licenciamento")
    add_bullet("O que faz: ", "Permite ao administrador gerenciar usuários, alterar seus planos (Basic, Pro, Premium), definir limites quantitativos e ativar Feature Flags Globais (como o Modo Copa do Mundo).")
    add_bullet("Como faz: ", "Painel SuperPage.tsx restrito via verificação de e-mail/função na tabela super_admins com integração ao useFeatureFlags.ts.")

    # ---------------------------------------------------------
    # 5. GUIA COMPLETO DE IMPLANTAÇÃO E MANUTENÇÃO
    # ---------------------------------------------------------
    doc.add_page_break()
    add_h1("5. Guia Completo de Implantação e Manutenção")

    add_p("A implantação do sistema Fluxo Financeiro em um novo ambiente ou servidor segue um fluxo padronizado de 6 fases:")

    add_h2("Fase 1: Preparação do Ambiente de Infraestrutura")
    add_bullet("Servidor Node.js: ", "Requerido Node.js v18.0.0+ ou Bun v1.x instalado no servidor de CI/CD.")
    add_bullet("Instância Supabase: ", "Criar um projeto na Supabase Cloud (ou instalar Supabase Self-Hosted via Docker). Anotar a SUPABASE_URL e a SUPABASE_ANON_KEY.")

    add_h2("Fase 2: Instalação e Configuração do Código Fonte")
    add_code_block(
        "# 1. Clonar o repositório\n"
        "git clone <URL_DO_REPOSITORIO>\n"
        "cd fluxo-financeiro\n\n"
        "# 2. Instalar dependências\n"
        "npm install\n\n"
        "# 3. Criar o arquivo de variáveis de ambiente .env\n"
        "VITE_SUPABASE_URL=https://seu-projeto.supabase.co\n"
        "VITE_SUPABASE_ANON_KEY=sua-chave-anonima-aqui"
    )

    add_h2("Fase 3: Implantação do Banco de Dados e Migrações SQL")
    add_p("As migrações do banco de dados estão localizadas na pasta ", "supabase/migrations/. ")
    p_mig = doc.paragraphs[-1]
    p_mig.add_run("Elas devem ser executadas em ordem numérica utilizando a Supabase CLI ou a interface SQL Editor do Supabase:")
    add_code_block(
        "# Executar migrações via Supabase CLI\n"
        "npx supabase db push"
    )

    add_h2("Fase 4: Deploy da Aplicação Frontend (Vercel / Cloudflare Pages / Nginx)")
    add_p("O projeto já possui configuração nativa no arquivo vercel.json para hospedagem na Vercel com roteamento SPA (Single Page Application):")
    add_code_block(
        "# Build de Produção\n"
        "npm run build\n\n"
        "# O resultado compilado será gerado na pasta ./dist ready for deployment."
    )

    add_h2("Fase 5: Implantação de Notificações PWA e Edge Functions")
    add_p("Para ativar as notificações push do PWA:")
    add_bullet("1. ", "Gerar o par de chaves VAPID usando npx web-push generate-vapid-keys.")
    add_bullet("2. ", "Fazer deploy da Edge Function send-push localizada em supabase/functions/send-push/.")

    add_h2("Fase 6: Validação de Qualidade e Suíte de Testes")
    add_p("Antes de liberar o sistema para uso comercial, execute os comandos de verificação automatizada:")
    add_code_block(
        "# Rodar validação completa (Encoding UTF-8 + Linter + Testes Unitários + Build)\n"
        "npm run validate"
    )

    # ---------------------------------------------------------
    # 6. GUIA DE TRANSFERÊNCIA DE PROPRIEDADE INTELECTUAL (HANDOFF)
    # ---------------------------------------------------------
    add_h1("6. Guia de Transferência de Propriedade Intelectual (Handoff)")

    add_p("Caso o aplicativo seja negociado ou vendido para um novo proprietário/cliente, este capítulo estabelece o protocolo de transferência segura de ativos:")

    add_h2("6.1. Checklist de Ativos a Serem Entregues")
    add_styled_table(
        ["Ativo de TI", "Localização / Repositório", "Ação de Transferência"],
        [
            ["Código Fonte Completo", "Repositório Git (GitHub / GitLab)", "Transferência de propriedade do repositório ou exportação de zip contendo histórico de commits."],
            ["Projeto Supabase", "Supabase Cloud Dashboard", "Transferência de ownership da organização Supabase para o e-mail do comprador."],
            ["Hospedagem Frontend", "Vercel / Cloudflare", "Migração do projeto de hospedagem e reconfiguração de domínio personalizado (DNS)."],
            ["Credenciais Super Admin", "Tabela super_admins", "Cadastro do e-mail do comprador na tabela de administradores master."],
            ["Documentação & Manuais", "Arquivos .docx e .md no repositório", "Entrega da documentação técnica e manual de cálculo do Score."]
        ],
        col_widths=[1.8, 2.0, 2.7]
    )

    add_h2("6.2. Matriz de Planos Sugerida para Comercialização SaaS")
    add_styled_table(
        ["Recurso / Módulo", "Plano Basic", "Plano Pro", "Plano Premium"],
        [
            ["Limite de Contas Bancárias", "Até 3 contas", "Até 10 contas", "Ilimitado"],
            ["Limite de Cartões de Crédito", "Até 2 cartões", "Até 5 cartões", "Ilimitado"],
            ["Relatórios Gerenciais PDF", "Básico", "Completo", "Completo + Insights IA"],
            ["Módulo de Dívidas & Acordos", "Não incluído", "Incluído", "Incluído"],
            ["Fluxo Score Algorítmico", "Não incluído", "Incluído", "Incluído"],
            ["Notificações Push PWA", "Não incluído", "Incluído", "Incluído"],
            ["Suporte Multi-Moeda / Temas", "Não incluído", "Não incluído", "Incluído"]
        ],
        col_widths=[2.2, 1.4, 1.4, 1.5]
    )

    add_callout(
        "Conclusão do Analista de Sistemas",
        "O Fluxo Financeiro representa uma solução pronta, robusta, altamente escalável e de fácil manutenção. Suas escolhas arquiteturais (React + TypeScript + Supabase BaaS + Tailwind CSS + PWA) garantem baixíssimo custo de infraestrutura operacional e máxima velocidade para futuras expansões comerciais.",
        "info"
    )

    # ---------------------------------------------------------
    # SAVE DOCUMENT
    # ---------------------------------------------------------
    output_filename = "Documentacao_Tecnica_e_Arquitetural_Fluxo.docx"
    output_path = os.path.abspath(output_filename)
    doc.save(output_path)
    print(f"DOCUMENT_CREATED_SUCCESSFULLY: {output_path}")

if __name__ == "__main__":
    create_document()
